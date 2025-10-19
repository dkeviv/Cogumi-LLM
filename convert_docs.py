import docx
import os

def convert_docx_to_markdown(docx_file, markdown_file):
    """Convert a Word document to Markdown format."""
    doc = docx.Document(docx_file)
    
    with open(markdown_file, 'w', encoding='utf-8') as f:
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if not text:
                f.write('\n')
                continue
            
            # Handle headings
            if paragraph.style.name.startswith('Heading'):
                level = paragraph.style.name.replace('Heading ', '').strip()
                if level.isdigit():
                    f.write('#' * int(level) + ' ' + text + '\n\n')
                else:
                    f.write('# ' + text + '\n\n')
            else:
                f.write(text + '\n\n')
        
        # Handle tables
        for table in doc.tables:
            f.write('\n')
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                f.write('| ' + ' | '.join(cells) + ' |\n')
                if i == 0:  # Add separator after header row
                    f.write('| ' + ' | '.join(['---'] * len(cells)) + ' |\n')
            f.write('\n')

# Convert the dev documents
dev_dir = '/Users/vivekdurairaj/Projects/Cogumi-LLM/docs/dev'
files = [
    'For Dev Final MVP and Phase 2 pipleine.docx',
    'For Dev_ COMPLETE TECHNICAL METHODOLOGY_ REVISED PIPELINE.docx'
]

for filename in files:
    docx_path = os.path.join(dev_dir, filename)
    md_filename = filename.replace('.docx', '.md')
    md_path = os.path.join(dev_dir, md_filename)
    
    print(f"Converting {filename} to {md_filename}...")
    convert_docx_to_markdown(docx_path, md_path)
    print(f"Done!")

print("\nAll documents converted successfully!")
